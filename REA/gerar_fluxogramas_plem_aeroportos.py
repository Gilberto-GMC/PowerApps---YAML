#!/usr/bin/env python3
"""Gera os fluxogramas PLEM de todos os aeroportos em DOCX estável.

O layout de cada página é rasterizado em alta resolução antes de ser inserido
no Word. Isso evita as quebras de posicionamento observadas com formas VML no
Word Desktop/Online e mantém o mesmo resultado em visualização e impressão.

Fontes de dados:
- tbl_Emergencias.xlsx
- tbl_FluxogramaAcionamentos.xlsx
- tbl_equipamentosAcionamentos.xlsx (IATA, quando disponível)

Modelo visual:
- páginas 1 a 8 de "Anexo 3 a 13 - Fluxogramas Acionamentos_PLEM SLZ.pdf"
"""

from __future__ import annotations

from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import date
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile
import json
import re
import unicodedata

import pymupdf
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent
PDF_MODELO = BASE_DIR / "Anexo 3 a 13 - Fluxogramas Acionamentos_PLEM SLZ.pdf"
ARQ_EMERGENCIAS = BASE_DIR / "tbl_Emergencias.xlsx"
ARQ_FLUXOS = BASE_DIR / "tbl_FluxogramaAcionamentos.xlsx"
ARQ_EQUIPAMENTOS = BASE_DIR / "tbl_equipamentosAcionamentos.xlsx"

PASTA_SAIDA = BASE_DIR / "Fluxogramas_PLEM_Aeroportos"
ARQ_CONSOLIDADO = PASTA_SAIDA / "00_Fluxogramas_PLEM_Todos_Aeroportos.docx"
ARQ_RELATORIO = PASTA_SAIDA / "00_Comparativo_Fluxos_Aeroportos.xlsx"
ARQ_MANIFESTO = PASTA_SAIDA / "00_Manifesto_Geracao.txt"
ARQ_PACOTE = BASE_DIR / "Fluxogramas_PLEM_Todos_Aeroportos.zip"

LARGURA_PT = 540.0
ALTURA_PT = 780.0
ESCALA = 4
LARGURA_PX = int(LARGURA_PT * ESCALA)
ALTURA_PX = int(ALTURA_PT * ESCALA)

MOLDURA_X = 37.12
MOLDURA_Y = 77.16
MOLDURA_LARGURA = 465.76
MOLDURA_ALTURA = 675.93
FAIXA_TITULO_ALTURA = 26.71

COR_TEXTO = "#1E1E64"
COR_INICIAL = "#FCEBD1"
COR_ACIONAR = "#FF0000"
COR_SOBREAVISO = "#FFC000"
COR_INFORMAR = "#2E75B6"
COR_CINZA = "#D9D9D9"
COR_GRADE = "#BFBFBF"
COR_PRETO = "#000000"
COR_BRANCO = "#FFFFFF"

FONTE_REGULAR = Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf")
FONTE_NEGRITO = Path("/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf")
if not FONTE_REGULAR.exists():
    FONTE_REGULAR = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
if not FONTE_NEGRITO.exists():
    FONTE_NEGRITO = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")


@dataclass
class Fluxo:
    aeroporto: str
    bloco: str
    iata: str
    emergencia_id: str
    titulo_origem: str
    titulo_apresentacao: str
    categoria: str
    grupos: dict[str, list[dict[str, object]]]
    assinatura: str

    @property
    def quantidade_nos(self) -> int:
        return sum(len(itens) for itens in self.grupos.values())


def px(valor_pt: float) -> int:
    return int(round(valor_pt * ESCALA))


def normalizar_chave(valor: object) -> str:
    texto = "" if valor is None else str(valor)
    texto = " ".join(texto.strip().upper().split())
    return "".join(
        caractere
        for caractere in unicodedata.normalize("NFKD", texto)
        if not unicodedata.combining(caractere)
    )


def normalizar_texto(valor: object) -> str:
    return re.sub(r"\s+", " ", "" if valor is None else str(valor)).strip().upper()


def id_texto(valor: object) -> str:
    if valor is None:
        return ""
    if isinstance(valor, float) and valor.is_integer():
        return str(int(valor))
    return str(valor).strip()


def numero_ordem(valor: object) -> int:
    try:
        return int(float(str(valor).strip()))
    except (TypeError, ValueError):
        return 999999


def ler_xlsx(caminho: Path) -> list[dict[str, object]]:
    workbook = load_workbook(caminho, read_only=True, data_only=True)
    planilha = workbook.active
    linhas = planilha.iter_rows(values_only=True)
    cabecalho = [str(valor).strip() if valor is not None else "" for valor in next(linhas)]
    return [
        dict(zip(cabecalho, linha))
        for linha in linhas
        if any(valor is not None for valor in linha)
    ]


def categoria_titulo(titulo: object) -> str:
    chave = normalizar_chave(titulo)
    if "PANPAN" in chave:
        return "PANPAN"
    if "MAYDAY" in chave:
        return "MAYDAY"
    if "FORA DA AREA DE ATUACAO DO SESCINC" in chave:
        return "FORA_AREA_SESCINC"
    if "MEDIC" in chave and "GERAL" in chave:
        return "MEDICAS"
    if "ARTIGOS PERIGOSOS" in chave:
        return "ARTIGOS_PERIGOSOS"
    if "DESASTRE NATURAL" in chave:
        return "DESASTRE_NATURAL"
    if "INCENDIO" in chave and "VEGETACAO" in chave:
        return "INCENDIO_VEGETACAO"
    if "INCENDIO" in chave and "INSTALACOES" in chave:
        return "INCENDIO_INSTALACOES"
    if "SALVAMENTO AQUATICO" in chave or "SESAQ" in chave:
        return "SESAQ"
    return "OUTRO"


ORDEM_CATEGORIAS = {
    "PANPAN": 1,
    "MAYDAY": 2,
    "FORA_AREA_SESCINC": 3,
    "MEDICAS": 4,
    "ARTIGOS_PERIGOSOS": 5,
    "DESASTRE_NATURAL": 6,
    "INCENDIO_VEGETACAO": 7,
    "INCENDIO_INSTALACOES": 8,
    "SESAQ": 9,
    "OUTRO": 99,
}

TITULOS_PADRONIZADOS = {
    "PANPAN": "Condição de Urgência - PANPAN",
    "MAYDAY": "Condição de Socorro - MAYDAY",
    "FORA_AREA_SESCINC": "Condição de Urgência e Socorro fora da área de atuação do SESCINC",
    "MEDICAS": "Emergências Médicas em Geral",
    "ARTIGOS_PERIGOSOS": "Emergências com Artigos Perigosos",
    "DESASTRE_NATURAL": "Emergências por Desastre Natural",
    "INCENDIO_VEGETACAO": "Emergências por Incêndio em Vegetação",
    "INCENDIO_INSTALACOES": "Emergências por Incêndio em Instalações",
    "SESAQ": "Serviço de Salvamento Aquático - SESAQ",
}


def titulo_apresentacao(titulo: object, categoria: str) -> str:
    if categoria in TITULOS_PADRONIZADOS:
        return TITULOS_PADRONIZADOS[categoria]
    texto = re.sub(r"\s+", " ", str(titulo or "")).strip()
    return texto if texto else "Fluxograma PLEM"


def assinatura_grupos(grupos: dict[str, list[dict[str, object]]]) -> str:
    estrutura = {
        nivel: [normalizar_texto(item.get("Titulo")) for item in grupos[nivel]]
        for nivel in ("ACIONAR", "SOBREAVISO", "INFORMAR")
    }
    serializado = json.dumps(estrutura, ensure_ascii=False, sort_keys=True)
    return sha256(serializado.encode("utf-8")).hexdigest()[:16]


def carregar_iatas(equipamentos: list[dict[str, object]]) -> dict[str, str]:
    contagem: dict[str, Counter[str]] = defaultdict(Counter)
    for item in equipamentos:
        aeroporto = str(item.get("Aeroporto") or "").strip()
        iata = str(item.get("IATA") or "").strip().upper()
        if aeroporto and iata:
            contagem[aeroporto][iata] += 1
    return {
        aeroporto: valores.most_common(1)[0][0]
        for aeroporto, valores in contagem.items()
    }


def carregar_dados() -> tuple[
    dict[str, list[Fluxo]],
    list[dict[str, object]],
    list[dict[str, object]],
    list[str],
]:
    emergencias = ler_xlsx(ARQ_EMERGENCIAS)
    itens_fluxo = ler_xlsx(ARQ_FLUXOS)
    equipamentos = ler_xlsx(ARQ_EQUIPAMENTOS)
    iatas = carregar_iatas(equipamentos)
    emergencias_por_id = {id_texto(item.get("ID")): item for item in emergencias}

    registros_por_emergencia: dict[str, list[dict[str, object]]] = defaultdict(list)
    orfaos: list[dict[str, object]] = []
    for item in itens_fluxo:
        emergencia_id = id_texto(item.get("ID_emergencia"))
        if emergencia_id not in emergencias_por_id:
            orfaos.append(item)
        else:
            registros_por_emergencia[emergencia_id].append(item)

    por_aeroporto: dict[str, list[Fluxo]] = defaultdict(list)
    sem_fluxo: list[dict[str, object]] = []
    avisos: list[str] = []

    for emergencia in emergencias:
        if normalizar_chave(emergencia.get("Acionamento")) != "PLEM":
            continue
        emergencia_id = id_texto(emergencia.get("ID"))
        aeroporto = str(emergencia.get("Aeroporto") or "").strip()
        registros = [
            item
            for item in registros_por_emergencia.get(emergencia_id, [])
            if str(item.get("Aeroporto") or "").strip() == aeroporto
        ]
        if not registros:
            sem_fluxo.append(emergencia)
            continue

        grupos: dict[str, list[dict[str, object]]] = {}
        for nivel in ("ACIONAR", "SOBREAVISO", "INFORMAR"):
            itens = [
                item
                for item in registros
                if normalizar_chave(item.get("Nível")) == nivel
            ]
            itens.sort(
                key=lambda item: (
                    numero_ordem(item.get("Ordem de prioridade")),
                    numero_ordem(item.get("ID")),
                )
            )
            prioridades = [numero_ordem(item.get("Ordem de prioridade")) for item in itens]
            repetidas = sorted(valor for valor, qtd in Counter(prioridades).items() if qtd > 1)
            if repetidas:
                avisos.append(
                    f"{aeroporto} / ID {emergencia_id} / {nivel}: "
                    f"prioridades repetidas {repetidas}; desempate pelo ID."
                )
            grupos[nivel] = itens

        categoria = categoria_titulo(emergencia.get("Titulo"))
        fluxo = Fluxo(
            aeroporto=aeroporto,
            bloco=str(emergencia.get("Bloco") or "").strip(),
            iata=iatas.get(aeroporto, ""),
            emergencia_id=emergencia_id,
            titulo_origem=str(emergencia.get("Titulo") or "").strip(),
            titulo_apresentacao=titulo_apresentacao(emergencia.get("Titulo"), categoria),
            categoria=categoria,
            grupos=grupos,
            assinatura=assinatura_grupos(grupos),
        )
        por_aeroporto[aeroporto].append(fluxo)

    for fluxos in por_aeroporto.values():
        fluxos.sort(
            key=lambda fluxo: (
                ORDEM_CATEGORIAS.get(fluxo.categoria, 99),
                numero_ordem(fluxo.emergencia_id),
            )
        )

    quantidade_fluxos = sum(len(fluxos) for fluxos in por_aeroporto.values())
    quantidade_nos = sum(
        fluxo.quantidade_nos for fluxos in por_aeroporto.values() for fluxo in fluxos
    )
    if len(por_aeroporto) != 16:
        raise RuntimeError(f"Esperados 16 aeroportos; encontrados {len(por_aeroporto)}.")
    if quantidade_fluxos != 129:
        raise RuntimeError(f"Esperados 129 fluxos; encontrados {quantidade_fluxos}.")
    if quantidade_nos != 1976:
        raise RuntimeError(f"Esperados 1.976 nós; encontrados {quantidade_nos}.")
    return dict(sorted(por_aeroporto.items())), sem_fluxo, orfaos, avisos


_fontes: dict[tuple[int, bool], ImageFont.FreeTypeFont] = {}


def fonte(tamanho_pt: float, negrito: bool = True) -> ImageFont.FreeTypeFont:
    tamanho_px = max(4, int(round(tamanho_pt * ESCALA)))
    chave = (tamanho_px, negrito)
    if chave not in _fontes:
        _fontes[chave] = ImageFont.truetype(
            str(FONTE_NEGRITO if negrito else FONTE_REGULAR), tamanho_px
        )
    return _fontes[chave]


def caixa_px(x: float, y: float, largura: float, altura: float) -> tuple[int, int, int, int]:
    return (px(x), px(y), px(x + largura), px(y + altura))


def quebrar_texto(
    desenho: ImageDraw.ImageDraw,
    texto: str,
    font: ImageFont.FreeTypeFont,
    largura_max: int,
) -> list[str]:
    palavras = texto.split()
    if not palavras:
        return [""]
    linhas: list[str] = []
    atual = palavras[0]
    for palavra in palavras[1:]:
        candidato = f"{atual} {palavra}"
        largura = desenho.textbbox((0, 0), candidato, font=font)[2]
        if largura <= largura_max:
            atual = candidato
        else:
            linhas.append(atual)
            atual = palavra
    linhas.append(atual)

    finais: list[str] = []
    for linha in linhas:
        if desenho.textbbox((0, 0), linha, font=font)[2] <= largura_max:
            finais.append(linha)
            continue
        fragmento = ""
        for caractere in linha:
            candidato = fragmento + caractere
            if fragmento and desenho.textbbox((0, 0), candidato, font=font)[2] > largura_max:
                finais.append(fragmento)
                fragmento = caractere
            else:
                fragmento = candidato
        if fragmento:
            finais.append(fragmento)
    return finais


def desenhar_texto_centralizado(
    desenho: ImageDraw.ImageDraw,
    bbox: tuple[int, int, int, int],
    texto: str,
    *,
    tamanho_max: float,
    tamanho_min: float,
    cor: str,
    negrito: bool = True,
    margem_pt: float = 2.0,
) -> None:
    x1, y1, x2, y2 = bbox
    largura_max = max(1, x2 - x1 - 2 * px(margem_pt))
    altura_max = max(1, y2 - y1 - 2 * px(margem_pt))
    tamanho_px_max = int(round(tamanho_max * ESCALA))
    tamanho_px_min = int(round(tamanho_min * ESCALA))
    escolhido = fonte(tamanho_min, negrito)
    linhas_escolhidas = [texto]
    altura_linha = escolhido.getbbox("Ag")[3] - escolhido.getbbox("Ag")[1]

    for tamanho_px in range(tamanho_px_max, tamanho_px_min - 1, -1):
        tamanho_pt = tamanho_px / ESCALA
        atual = fonte(tamanho_pt, negrito)
        linhas = quebrar_texto(desenho, texto, atual, largura_max)
        medidas = atual.getbbox("Ag")
        altura = medidas[3] - medidas[1]
        espacamento = max(0, int(round(altura * 0.08)))
        altura_total = len(linhas) * altura + max(0, len(linhas) - 1) * espacamento
        if altura_total <= altura_max:
            escolhido = atual
            linhas_escolhidas = linhas
            altura_linha = altura
            break

    espacamento = max(0, int(round(altura_linha * 0.08)))
    altura_total = len(linhas_escolhidas) * altura_linha + max(0, len(linhas_escolhidas) - 1) * espacamento
    y = y1 + (y2 - y1 - altura_total) / 2
    for linha in linhas_escolhidas:
        medidas = desenho.textbbox((0, 0), linha, font=escolhido)
        largura = medidas[2] - medidas[0]
        x = x1 + (x2 - x1 - largura) / 2
        desenho.text((round(x), round(y - medidas[1])), linha, font=escolhido, fill=cor)
        y += altura_linha + espacamento


def desenhar_retangulo(
    desenho: ImageDraw.ImageDraw,
    x: float,
    y: float,
    largura: float,
    altura: float,
    *,
    contorno: str | None = None,
    preenchimento: str | None = None,
    espessura: float = 1.0,
) -> None:
    desenho.rectangle(
        caixa_px(x, y, largura, altura),
        fill=preenchimento,
        outline=contorno,
        width=max(1, px(espessura)) if contorno else 1,
    )


def linha_tracejada(
    desenho: ImageDraw.ImageDraw,
    inicio: tuple[float, float],
    fim: tuple[float, float],
    *,
    cor: str,
    espessura: float,
    traco: float = 8.0,
    intervalo: float = 5.0,
) -> None:
    x1, y1 = inicio
    x2, y2 = fim
    largura = max(1, px(espessura))
    if abs(y2 - y1) < 0.01:
        atual = x1
        while atual < x2:
            final = min(x2, atual + traco)
            desenho.line((px(atual), px(y1), px(final), px(y2)), fill=cor, width=largura)
            atual += traco + intervalo
    else:
        atual = y1
        while atual < y2:
            final = min(y2, atual + traco)
            desenho.line((px(x1), px(atual), px(x2), px(final)), fill=cor, width=largura)
            atual += traco + intervalo


def retangulo_tracejado(
    desenho: ImageDraw.ImageDraw,
    x: float,
    y: float,
    largura: float,
    altura: float,
    *,
    cor: str,
    espessura: float,
) -> None:
    linha_tracejada(desenho, (x, y), (x + largura, y), cor=cor, espessura=espessura)
    linha_tracejada(
        desenho, (x, y + altura), (x + largura, y + altura), cor=cor, espessura=espessura
    )
    linha_tracejada(desenho, (x, y), (x, y + altura), cor=cor, espessura=espessura)
    linha_tracejada(
        desenho, (x + largura, y), (x + largura, y + altura), cor=cor, espessura=espessura
    )


def desenhar_seta_vertical(
    desenho: ImageDraw.ImageDraw,
    x: float,
    y_inicio: float,
    y_fim: float,
) -> None:
    if y_fim <= y_inicio:
        return
    ponta = min(3.2, max(1.8, (y_fim - y_inicio) * 0.35))
    desenho.line((px(x), px(y_inicio), px(x), px(y_fim - ponta)), fill=COR_PRETO, width=px(0.8))
    desenho.polygon(
        (
            (px(x - ponta), px(y_fim - ponta)),
            (px(x + ponta), px(y_fim - ponta)),
            (px(x), px(y_fim)),
        ),
        fill=COR_PRETO,
    )


def condicional(texto: str) -> bool:
    chave = normalizar_chave(texto)
    return bool(
        re.search(r"\(\s*(SE|QUANDO|CONFORME)", chave)
        or " SE " in chave
        or " QUANDO " in chave
        or " CONFORME " in chave
    )


def extrair_logo() -> Image.Image:
    pdf = pymupdf.open(PDF_MODELO)
    imagem = pdf[0].get_images(full=True)[0]
    base = pymupdf.Pixmap(pdf, imagem[0])
    mascara = pymupdf.Pixmap(pdf, imagem[1])
    pixmap = pymupdf.Pixmap(base, mascara)
    return Image.open(BytesIO(pixmap.tobytes("png"))).convert("RGBA")


def desenhar_cabecalho(
    imagem: Image.Image,
    desenho: ImageDraw.ImageDraw,
    logo: Image.Image,
    fluxo: Fluxo,
    indice: int,
    total: int,
) -> None:
    x = 37.12
    y = 27.0
    linha = 13.0
    larguras = (90.0, 317.76, 58.0)
    x_centro = x + larguras[0]
    x_direita = x_centro + larguras[1]
    celulas = (
        (x, y, larguras[0], linha * 2),
        (x_centro, y, larguras[1], linha),
        (x_direita, y, larguras[2], linha),
        (x_centro, y + linha, larguras[1], linha),
        (x_direita, y + linha, larguras[2], linha * 2),
        (x, y + linha * 2, larguras[0], linha),
        (x_centro, y + linha * 2, larguras[1], linha),
    )
    for cx, cy, largura, altura in celulas:
        desenhar_retangulo(
            desenho,
            cx,
            cy,
            largura,
            altura,
            contorno=COR_GRADE,
            espessura=0.55,
        )

    logo_largura = px(56.4)
    logo_altura = round(logo.height * logo_largura / logo.width)
    logo_redimensionado = logo.resize((logo_largura, logo_altura), Image.Resampling.LANCZOS)
    imagem.alpha_composite(logo_redimensionado, (px(53.9), px(27.0)))

    aeroporto_cabecalho = f"AEROPORTO: {fluxo.aeroporto}"
    if fluxo.iata:
        aeroporto_cabecalho += f" ({fluxo.iata})"
    desenhar_texto_centralizado(
        desenho,
        caixa_px(x_centro, y, larguras[1], linha),
        aeroporto_cabecalho,
        tamanho_max=6.5,
        tamanho_min=5.0,
        cor=COR_PRETO,
        negrito=False,
        margem_pt=0.5,
    )
    desenhar_texto_centralizado(
        desenho,
        caixa_px(x_direita, y, larguras[2], linha),
        "PLEM",
        tamanho_max=6.5,
        tamanho_min=5.0,
        cor=COR_PRETO,
        negrito=True,
        margem_pt=0.3,
    )
    desenhar_texto_centralizado(
        desenho,
        caixa_px(x_centro, y + linha, larguras[1], linha),
        f"ID DA EMERGÊNCIA: {fluxo.emergencia_id}",
        tamanho_max=6.3,
        tamanho_min=5.0,
        cor=COR_PRETO,
        negrito=False,
        margem_pt=0.3,
    )
    desenhar_texto_centralizado(
        desenho,
        caixa_px(x_direita, y + linha, larguras[2], linha * 2),
        f"{fluxo.bloco}\n{date.today().strftime('%d/%m/%Y')}",
        tamanho_max=5.4,
        tamanho_min=4.3,
        cor=COR_PRETO,
        negrito=False,
        margem_pt=0.3,
    )
    desenhar_texto_centralizado(
        desenho,
        caixa_px(x, y + linha * 2, larguras[0], linha),
        f"FLUXO {indice}/{total}",
        tamanho_max=6.3,
        tamanho_min=5.0,
        cor=COR_PRETO,
        negrito=False,
        margem_pt=0.3,
    )
    desenhar_texto_centralizado(
        desenho,
        caixa_px(x_centro, y + linha * 2, larguras[1], linha),
        f"Fluxograma para {fluxo.titulo_apresentacao}",
        tamanho_max=5.8,
        tamanho_min=4.4,
        cor=COR_PRETO,
        negrito=True,
        margem_pt=0.3,
    )


def perfil_topo(categoria: str) -> tuple[tuple[str, ...], bool, float]:
    if categoria in {"PANPAN", "MAYDAY"}:
        return ("ATS", "APOC (2º)", "SUPERVISOR", "COE/AGENTE"), True, 268.0
    if categoria == "FORA_AREA_SESCINC":
        return ("ATS", "APOC", "SUPERVISOR", "COE/AGENTE"), False, 268.0
    if categoria in {"MEDICAS", "ARTIGOS_PERIGOSOS"}:
        return ("ATS/IDENTIFICADOR DA OCORRÊNCIA", "APOC"), False, 202.0
    return ("IDENTIFICADOR DA OCORRÊNCIA", "APOC/COE"), False, 208.0


def desenhar_topo(
    desenho: ImageDraw.ImageDraw,
    categoria: str,
) -> tuple[float, float]:
    nos, lateral, y_ramificacao = perfil_topo(categoria)
    centro_x = 270.0
    largura = 95.0 if len(nos) == 4 else 112.0
    altura = 25.0 if len(nos) == 4 else 28.0
    intervalo = 8.0 if len(nos) == 4 else 11.0
    y = 116.0

    posicoes: list[tuple[float, float]] = []
    for indice, no in enumerate(nos):
        x = centro_x - largura / 2
        desenhar_retangulo(
            desenho,
            x,
            y,
            largura,
            altura,
            contorno=COR_TEXTO,
            preenchimento=COR_INICIAL if indice == 0 else COR_BRANCO,
            espessura=1.25,
        )
        desenhar_texto_centralizado(
            desenho,
            caixa_px(x, y, largura, altura),
            no,
            tamanho_max=8.5,
            tamanho_min=5.5,
            cor=COR_TEXTO,
            negrito=True,
            margem_pt=1.2,
        )
        posicoes.append((y, y + altura))
        if indice < len(nos) - 1:
            desenhar_seta_vertical(desenho, centro_x, y + altura, y + altura + intervalo)
        y += altura + intervalo

    if lateral:
        lateral_x = 58.0
        lateral_y = posicoes[1][0]
        largura_lateral = 127.56
        desenhar_retangulo(
            desenho,
            lateral_x,
            lateral_y,
            largura_lateral,
            altura,
            contorno=COR_ACIONAR,
            preenchimento=COR_BRANCO,
            espessura=2.25,
        )
        desenhar_texto_centralizado(
            desenho,
            caixa_px(lateral_x, lateral_y, largura_lateral, altura),
            "SESCINC (1º)",
            tamanho_max=8.0,
            tamanho_min=6.0,
            cor=COR_TEXTO,
            negrito=True,
            margem_pt=1.0,
        )
        linha_y = lateral_y + altura / 2
        desenho.line(
            (
                px(lateral_x + largura_lateral),
                px(linha_y),
                px(centro_x - largura / 2),
                px(linha_y),
            ),
            fill=COR_PRETO,
            width=px(0.8),
        )
        desenhar_texto_centralizado(
            desenho,
            caixa_px(188.0, linha_y - 13.0, 34.0, 12.0),
            "Confirma",
            tamanho_max=6.3,
            tamanho_min=5.0,
            cor=COR_PRETO,
            negrito=False,
            margem_pt=0,
        )

    fim_topo = posicoes[-1][1]
    desenho.line(
        (px(centro_x), px(fim_topo), px(centro_x), px(y_ramificacao)),
        fill=COR_PRETO,
        width=px(0.8),
    )
    return y_ramificacao, fim_topo


def desenhar_grupos(
    desenho: ImageDraw.ImageDraw,
    grupos: dict[str, list[dict[str, object]]],
    y_ramificacao: float,
) -> None:
    definicoes = (
        ("ACIONAR", 58.0, COR_ACIONAR),
        ("SOBREAVISO", 207.0, COR_SOBREAVISO),
        ("INFORMAR", 356.0, COR_INFORMAR),
    )
    largura_caixa = 127.56
    largura_rotulo = 68.03
    altura_rotulo = 17.01
    y_rotulo = y_ramificacao + 8.0
    y_primeiro = y_rotulo + altura_rotulo + 17.0
    limite_inferior = MOLDURA_Y + MOLDURA_ALTURA - 10.0
    espaco_seta = 7.0
    quantidade_maxima = max(len(grupos[nivel]) for nivel, _, _ in definicoes)
    disponivel = limite_inferior - y_primeiro
    altura_caixa = min(
        42.52,
        (disponivel - espaco_seta * max(0, quantidade_maxima - 1)) / max(1, quantidade_maxima),
    )
    if altura_caixa < 25.0:
        raise RuntimeError(
            f"Layout insuficiente para {quantidade_maxima} itens na mesma coluna."
        )

    centros = [x + largura_caixa / 2 for _, x, _ in definicoes]
    desenho.line(
        (px(centros[0]), px(y_ramificacao), px(centros[-1]), px(y_ramificacao)),
        fill=COR_PRETO,
        width=px(0.8),
    )

    for nivel, x, cor in definicoes:
        centro = x + largura_caixa / 2
        desenho.line(
            (px(centro), px(y_ramificacao), px(centro), px(y_rotulo)),
            fill=COR_PRETO,
            width=px(0.8),
        )
        x_rotulo = x + (largura_caixa - largura_rotulo) / 2
        desenhar_retangulo(
            desenho,
            x_rotulo,
            y_rotulo,
            largura_rotulo,
            altura_rotulo,
            preenchimento=cor,
        )
        desenhar_texto_centralizado(
            desenho,
            caixa_px(x_rotulo, y_rotulo, largura_rotulo, altura_rotulo),
            nivel,
            tamanho_max=6.0,
            tamanho_min=5.0,
            cor=COR_PRETO,
            negrito=True,
            margem_pt=0.4,
        )

        itens = grupos[nivel]
        if not itens:
            continue
        desenhar_seta_vertical(
            desenho,
            centro,
            y_rotulo + altura_rotulo,
            y_primeiro,
        )
        y = y_primeiro
        for indice, item in enumerate(itens):
            texto = normalizar_texto(item.get("Titulo"))
            if condicional(texto):
                retangulo_tracejado(
                    desenho,
                    x,
                    y,
                    largura_caixa,
                    altura_caixa,
                    cor=cor,
                    espessura=2.25,
                )
            else:
                desenhar_retangulo(
                    desenho,
                    x,
                    y,
                    largura_caixa,
                    altura_caixa,
                    contorno=cor,
                    preenchimento=COR_BRANCO,
                    espessura=2.25,
                )
            desenhar_texto_centralizado(
                desenho,
                caixa_px(x, y, largura_caixa, altura_caixa),
                texto,
                tamanho_max=9.0,
                tamanho_min=4.8,
                cor=COR_TEXTO,
                negrito=True,
                margem_pt=2.0,
            )
            if indice < len(itens) - 1:
                desenhar_seta_vertical(
                    desenho,
                    centro,
                    y + altura_caixa,
                    y + altura_caixa + espaco_seta,
                )
            y += altura_caixa + espaco_seta


def renderizar_fluxo(
    fluxo: Fluxo,
    indice: int,
    total: int,
    logo: Image.Image,
) -> bytes:
    imagem = Image.new("RGBA", (LARGURA_PX, ALTURA_PX), COR_BRANCO)
    desenho = ImageDraw.Draw(imagem)
    desenhar_cabecalho(imagem, desenho, logo, fluxo, indice, total)
    desenhar_retangulo(
        desenho,
        MOLDURA_X,
        MOLDURA_Y,
        MOLDURA_LARGURA,
        MOLDURA_ALTURA,
        contorno=COR_PRETO,
        preenchimento=COR_BRANCO,
        espessura=1.0,
    )
    desenhar_retangulo(
        desenho,
        MOLDURA_X,
        MOLDURA_Y,
        MOLDURA_LARGURA,
        FAIXA_TITULO_ALTURA,
        contorno=COR_PRETO,
        preenchimento=COR_CINZA,
        espessura=1.0,
    )
    desenhar_texto_centralizado(
        desenho,
        caixa_px(MOLDURA_X, MOLDURA_Y, MOLDURA_LARGURA, FAIXA_TITULO_ALTURA),
        fluxo.titulo_apresentacao,
        tamanho_max=14.0,
        tamanho_min=9.0,
        cor=COR_PRETO,
        negrito=True,
        margem_pt=2.0,
    )
    y_ramificacao, _ = desenhar_topo(desenho, fluxo.categoria)
    desenhar_grupos(desenho, fluxo.grupos, y_ramificacao)

    saida = BytesIO()
    imagem.convert("RGB").save(
        saida,
        format="PNG",
        dpi=(72 * ESCALA, 72 * ESCALA),
        optimize=True,
        compress_level=9,
    )
    return saida.getvalue()


def configurar_docx(documento: Document, titulo: str) -> None:
    secao = documento.sections[0]
    secao.page_width = Pt(LARGURA_PT)
    secao.page_height = Pt(ALTURA_PT)
    secao.top_margin = Pt(10)
    secao.bottom_margin = Pt(10)
    secao.left_margin = Pt(10)
    secao.right_margin = Pt(10)
    secao.header_distance = Pt(0)
    secao.footer_distance = Pt(0)
    documento.core_properties.title = titulo
    documento.core_properties.subject = "Fluxogramas PLEM gerados das listas SharePoint"
    documento.core_properties.author = "AirportNow / Motiva"
    documento.core_properties.keywords = "PLEM, fluxograma, acionamentos, Power Apps"


def adicionar_pagina_png(documento: Document, png: bytes, ultima: bool) -> None:
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    ppr = paragrafo._p.get_or_add_pPr()
    snap = OxmlElement("w:snapToGrid")
    snap.set(qn("w:val"), "0")
    ppr.append(snap)
    largura = 520.0
    altura = largura * ALTURA_PT / LARGURA_PT
    paragrafo.add_run().add_picture(BytesIO(png), width=Pt(largura), height=Pt(altura))
    if not ultima:
        paragrafo.add_run().add_break(WD_BREAK.PAGE)


def salvar_docx(caminho: Path, titulo: str, paginas: list[bytes]) -> None:
    documento = Document()
    configurar_docx(documento, titulo)
    for indice, pagina in enumerate(paginas):
        adicionar_pagina_png(documento, pagina, indice == len(paginas) - 1)
    documento.save(caminho)


def definir_largura_celula(celula, largura_pt: float) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(round(largura_pt * 20))))
    tc_w.set(qn("w:type"), "dxa")


def definir_margens_celula(celula, margem_pt: float = 0.5) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    valor = str(int(round(margem_pt * 20)))
    for lado in ("top", "start", "bottom", "end"):
        elemento = tc_mar.find(qn(f"w:{lado}"))
        if elemento is None:
            elemento = OxmlElement(f"w:{lado}")
            tc_mar.append(elemento)
        elemento.set(qn("w:w"), valor)
        elemento.set(qn("w:type"), "dxa")


def definir_bordas_celula(
    celula,
    *,
    cor: str,
    espessura_pt: float,
) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    bordas = tc_pr.first_child_found_in("w:tcBorders")
    if bordas is None:
        bordas = OxmlElement("w:tcBorders")
        tc_pr.append(bordas)
    for lado in ("top", "start", "bottom", "end"):
        borda = bordas.find(qn(f"w:{lado}"))
        if borda is None:
            borda = OxmlElement(f"w:{lado}")
            bordas.append(borda)
        borda.set(qn("w:val"), "single")
        borda.set(qn("w:sz"), str(max(2, int(round(espessura_pt * 8)))))
        borda.set(qn("w:space"), "0")
        borda.set(qn("w:color"), cor.lstrip("#"))


def definir_sombreamento_celula(celula, cor: str) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), cor.lstrip("#"))
    shd.set(qn("w:val"), "clear")


def formatar_texto_celula(
    celula,
    texto: str,
    *,
    tamanho_pt: float,
    negrito: bool,
) -> None:
    celula.text = ""
    celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragrafo = celula.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    paragrafo.paragraph_format.line_spacing = 1
    linhas = texto.split("\n")
    for indice, linha in enumerate(linhas):
        if indice:
            paragrafo.add_run().add_break()
        run = paragrafo.add_run(linha)
        run.bold = negrito
        run.font.name = "Arial"
        run.font.size = Pt(tamanho_pt)
        run.font.color.rgb = None


def definir_tabela_fixa(tabela, largura_pt: float) -> None:
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    tbl_pr = tabela._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(round(largura_pt * 20))))
    tbl_w.set(qn("w:type"), "dxa")


def definir_altura_linha(linha, altura_pt: float) -> None:
    linha.height = Pt(altura_pt)
    linha.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    tr_pr = linha._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def png_logo(logo: Image.Image) -> bytes:
    saida = BytesIO()
    logo.save(saida, format="PNG", optimize=True)
    return saida.getvalue()


def adicionar_cabecalho_editavel(
    documento: Document,
    fluxo: Fluxo,
    indice: int,
    total: int,
    logo_png: bytes,
) -> None:
    larguras = (90.0, 317.76, 58.0)
    tabela = documento.add_table(rows=3, cols=3)
    definir_tabela_fixa(tabela, MOLDURA_LARGURA)
    for linha in tabela.rows:
        definir_altura_linha(linha, 13.0)
        for coluna, celula in enumerate(linha.cells):
            definir_largura_celula(celula, larguras[coluna])
            definir_margens_celula(celula, 0.4)
            definir_bordas_celula(celula, cor=COR_GRADE, espessura_pt=0.55)

    celula_logo = tabela.cell(0, 0).merge(tabela.cell(1, 0))
    celula_data = tabela.cell(1, 2).merge(tabela.cell(2, 2))
    definir_largura_celula(celula_logo, larguras[0])
    definir_largura_celula(celula_data, larguras[2])
    definir_bordas_celula(celula_logo, cor=COR_GRADE, espessura_pt=0.55)
    definir_bordas_celula(celula_data, cor=COR_GRADE, espessura_pt=0.55)
    celula_logo.text = ""
    celula_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragrafo_logo = celula_logo.paragraphs[0]
    paragrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo_logo.paragraph_format.space_before = Pt(0)
    paragrafo_logo.paragraph_format.space_after = Pt(0)
    paragrafo_logo.add_run().add_picture(BytesIO(logo_png), width=Pt(56.4))

    aeroporto = f"AEROPORTO: {fluxo.aeroporto}"
    if fluxo.iata:
        aeroporto += f" ({fluxo.iata})"
    formatar_texto_celula(tabela.cell(0, 1), aeroporto, tamanho_pt=6.5, negrito=False)
    formatar_texto_celula(tabela.cell(0, 2), "PLEM", tamanho_pt=6.5, negrito=True)
    formatar_texto_celula(
        tabela.cell(1, 1),
        f"ID DA EMERGÊNCIA: {fluxo.emergencia_id}",
        tamanho_pt=6.3,
        negrito=False,
    )
    formatar_texto_celula(
        celula_data,
        f"{fluxo.bloco}\n{date.today().strftime('%d/%m/%Y')}",
        tamanho_pt=5.2,
        negrito=False,
    )
    formatar_texto_celula(
        tabela.cell(2, 0),
        f"FLUXO {indice}/{total}",
        tamanho_pt=6.3,
        negrito=False,
    )
    formatar_texto_celula(
        tabela.cell(2, 1),
        f"Fluxograma para {fluxo.titulo_apresentacao}",
        tamanho_pt=5.6 if len(fluxo.titulo_apresentacao) < 62 else 4.8,
        negrito=True,
    )


def adicionar_espacador(documento: Document, altura_pt: float) -> None:
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    paragrafo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragrafo.paragraph_format.line_spacing = Pt(altura_pt)
    run = paragrafo.add_run("\u00a0")
    run.font.size = Pt(1)


def adicionar_titulo_editavel(documento: Document, fluxo: Fluxo) -> None:
    tabela = documento.add_table(rows=1, cols=1)
    definir_tabela_fixa(tabela, MOLDURA_LARGURA)
    definir_altura_linha(tabela.rows[0], FAIXA_TITULO_ALTURA)
    celula = tabela.cell(0, 0)
    definir_largura_celula(celula, MOLDURA_LARGURA)
    definir_margens_celula(celula, 1.0)
    definir_bordas_celula(celula, cor=COR_PRETO, espessura_pt=1.0)
    definir_sombreamento_celula(celula, COR_CINZA)
    formatar_texto_celula(
        celula,
        fluxo.titulo_apresentacao,
        tamanho_pt=14.0 if len(fluxo.titulo_apresentacao) < 58 else 11.5,
        negrito=True,
    )


def recortar_corpo_fluxograma(pagina_png: bytes) -> bytes:
    pagina = Image.open(BytesIO(pagina_png)).convert("RGB")
    esquerda = px(MOLDURA_X)
    topo = px(MOLDURA_Y + FAIXA_TITULO_ALTURA)
    direita = px(MOLDURA_X + MOLDURA_LARGURA)
    inferior = px(MOLDURA_Y + MOLDURA_ALTURA)
    corpo = pagina.crop((esquerda, topo, direita, inferior))
    saida = BytesIO()
    corpo.save(saida, format="PNG", dpi=(72 * ESCALA, 72 * ESCALA), optimize=True, compress_level=9)
    return saida.getvalue()


def adicionar_corpo_fixo(
    documento: Document,
    pagina_png: bytes,
    *,
    ultima: bool,
) -> None:
    corpo_png = recortar_corpo_fluxograma(pagina_png)
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    ppr = paragrafo._p.get_or_add_pPr()
    snap = OxmlElement("w:snapToGrid")
    snap.set(qn("w:val"), "0")
    ppr.append(snap)
    paragrafo.add_run().add_picture(
        BytesIO(corpo_png),
        width=Pt(MOLDURA_LARGURA),
        height=Pt(646.0),
    )
    if not ultima:
        paragrafo.add_run().add_break(WD_BREAK.PAGE)


def salvar_consolidado_hibrido(
    caminho: Path,
    itens: list[tuple[Fluxo, int, int, bytes]],
    logo: Image.Image,
) -> None:
    documento = Document()
    secao = documento.sections[0]
    secao.page_width = Pt(LARGURA_PT)
    secao.page_height = Pt(ALTURA_PT)
    secao.top_margin = Pt(27.0)
    secao.bottom_margin = Pt(24.0)
    secao.left_margin = Pt(MOLDURA_X)
    secao.right_margin = Pt(LARGURA_PT - MOLDURA_X - MOLDURA_LARGURA)
    secao.header_distance = Pt(0)
    secao.footer_distance = Pt(0)
    documento.core_properties.title = "Fluxogramas PLEM — Todos os Aeroportos"
    documento.core_properties.subject = "Cabeçalhos editáveis e fluxogramas fixos"
    documento.core_properties.author = "AirportNow / Motiva"
    documento.core_properties.keywords = "PLEM, fluxograma, cabeçalho editável, Power Apps"
    logo_bytes = png_logo(logo)

    for posicao, (fluxo, indice, total, pagina_png) in enumerate(itens):
        adicionar_cabecalho_editavel(documento, fluxo, indice, total, logo_bytes)
        adicionar_espacador(documento, 10.0)
        adicionar_titulo_editavel(documento, fluxo)
        adicionar_corpo_fixo(documento, pagina_png, ultima=posicao == len(itens) - 1)
    documento.save(caminho)


def slug(texto: str) -> str:
    valor = normalizar_chave(texto)
    valor = re.sub(r"[^A-Z0-9]+", "_", valor).strip("_")
    return valor or "AEROPORTO"


def ajustar_larguras(planilha) -> None:
    for coluna in planilha.columns:
        letra = get_column_letter(coluna[0].column)
        maior = max(len(str(celula.value or "")) for celula in coluna)
        planilha.column_dimensions[letra].width = min(80, max(12, maior + 2))


def gerar_relatorio(
    por_aeroporto: dict[str, list[Fluxo]],
    sem_fluxo: list[dict[str, object]],
    orfaos: list[dict[str, object]],
    avisos: list[str],
) -> None:
    workbook = Workbook()
    resumo = workbook.active
    resumo.title = "Resumo"
    resumo.append(["Aeroporto", "IATA", "Bloco", "Fluxos", "Nós", "Assinatura geral"])
    for aeroporto, fluxos in por_aeroporto.items():
        assinatura_geral = sha256(
            "|".join(fluxo.assinatura for fluxo in fluxos).encode("utf-8")
        ).hexdigest()[:16]
        resumo.append(
            [
                aeroporto,
                fluxos[0].iata,
                fluxos[0].bloco,
                len(fluxos),
                sum(fluxo.quantidade_nos for fluxo in fluxos),
                assinatura_geral,
            ]
        )

    detalhes = workbook.create_sheet("Fluxos")
    detalhes.append(
        [
            "Aeroporto",
            "IATA",
            "ID emergência",
            "Categoria",
            "Título da lista",
            "Título padronizado",
            "Acionar",
            "Sobreaviso",
            "Informar",
            "Total",
            "Assinatura",
        ]
    )
    for aeroporto, fluxos in por_aeroporto.items():
        for fluxo in fluxos:
            detalhes.append(
                [
                    aeroporto,
                    fluxo.iata,
                    fluxo.emergencia_id,
                    fluxo.categoria,
                    fluxo.titulo_origem,
                    fluxo.titulo_apresentacao,
                    len(fluxo.grupos["ACIONAR"]),
                    len(fluxo.grupos["SOBREAVISO"]),
                    len(fluxo.grupos["INFORMAR"]),
                    fluxo.quantidade_nos,
                    fluxo.assinatura,
                ]
            )

    comparativo = workbook.create_sheet("Configurações distintas")
    comparativo.append(["Categoria", "Assinatura", "Aeroportos", "Quantidade"])
    agrupado: dict[tuple[str, str], list[str]] = defaultdict(list)
    for aeroporto, fluxos in por_aeroporto.items():
        for fluxo in fluxos:
            agrupado[(fluxo.categoria, fluxo.assinatura)].append(aeroporto)
    for (categoria, assinatura), aeroportos in sorted(agrupado.items()):
        comparativo.append([categoria, assinatura, ", ".join(aeroportos), len(aeroportos)])

    anomalias = workbook.create_sheet("Anomalias")
    anomalias.append(["Tipo", "Aeroporto/ID", "Detalhe"])
    for item in sem_fluxo:
        anomalias.append(
            [
                "Emergência PLEM sem nós",
                f"{item.get('Aeroporto')} / {id_texto(item.get('ID'))}",
                str(item.get("Titulo") or ""),
            ]
        )
    for item in orfaos:
        anomalias.append(
            [
                "Nó sem emergência correspondente",
                f"{item.get('Aeroporto')} / emergência {id_texto(item.get('ID_emergencia'))}",
                str(item.get("Titulo") or ""),
            ]
        )
    for aviso in avisos:
        anomalias.append(["Prioridade repetida", "", aviso])
    for aeroporto, fluxos in por_aeroporto.items():
        if not fluxos[0].iata:
            anomalias.append(["IATA ausente nas listas", aeroporto, "Cabeçalho gerado sem IATA."])

    fill = PatternFill("solid", fgColor="1E1E64")
    for planilha in workbook.worksheets:
        for celula in planilha[1]:
            celula.fill = fill
            celula.font = Font(color="FFFFFF", bold=True)
            celula.alignment = Alignment(horizontal="center", vertical="center")
        planilha.freeze_panes = "A2"
        planilha.auto_filter.ref = planilha.dimensions
        ajustar_larguras(planilha)
    workbook.save(ARQ_RELATORIO)


def gerar_manifesto(
    por_aeroporto: dict[str, list[Fluxo]],
    sem_fluxo: list[dict[str, object]],
    orfaos: list[dict[str, object]],
    avisos: list[str],
    arquivos: list[Path],
) -> None:
    linhas = [
        "GERAÇÃO DOS FLUXOGRAMAS PLEM — TODOS OS AEROPORTOS",
        "",
        f"Data: {date.today().strftime('%d/%m/%Y')}",
        f"Aeroportos: {len(por_aeroporto)}",
        f"Fluxogramas: {sum(len(v) for v in por_aeroporto.values())}",
        f"Nós exportados: {sum(f.quantidade_nos for v in por_aeroporto.values() for f in v)}",
        f"Emergências PLEM sem nós: {len(sem_fluxo)}",
        f"Nós órfãos: {len(orfaos)}",
        "",
        "ARQUIVOS POR AEROPORTO",
    ]
    for aeroporto, fluxos in por_aeroporto.items():
        linhas.append(
            f"- {aeroporto}: {len(fluxos)} fluxos, "
            f"{sum(fluxo.quantidade_nos for fluxo in fluxos)} nós"
        )
    linhas.extend(["", "AVISOS DE PRIORIDADE"])
    linhas.extend(f"- {aviso}" for aviso in avisos)
    linhas.extend(["", "ARQUIVOS GERADOS"])
    linhas.extend(f"- {arquivo.name}" for arquivo in arquivos)
    ARQ_MANIFESTO.write_text("\n".join(linhas) + "\n", encoding="utf-8")


def gerar_pacote(arquivos: Iterable[Path]) -> None:
    with ZipFile(ARQ_PACOTE, "w", compression=ZIP_DEFLATED, compresslevel=9) as pacote:
        for arquivo in arquivos:
            pacote.write(arquivo, arcname=arquivo.name)


def main() -> int:
    for caminho in (PDF_MODELO, ARQ_EMERGENCIAS, ARQ_FLUXOS, ARQ_EQUIPAMENTOS):
        if not caminho.exists():
            raise FileNotFoundError(f"Arquivo obrigatório ausente: {caminho}")

    PASTA_SAIDA.mkdir(parents=True, exist_ok=True)
    por_aeroporto, sem_fluxo, orfaos, avisos = carregar_dados()
    logo = extrair_logo()
    paginas_por_aeroporto: dict[str, list[bytes]] = {}
    arquivos_individuais: list[Path] = []
    itens_consolidados: list[tuple[Fluxo, int, int, bytes]] = []

    for aeroporto, fluxos in por_aeroporto.items():
        paginas = [
            renderizar_fluxo(fluxo, indice, len(fluxos), logo)
            for indice, fluxo in enumerate(fluxos, start=1)
        ]
        paginas_por_aeroporto[aeroporto] = paginas
        itens_consolidados.extend(
            (fluxo, indice, len(fluxos), pagina)
            for indice, (fluxo, pagina) in enumerate(zip(fluxos, paginas), start=1)
        )
        caminho = PASTA_SAIDA / f"Fluxogramas_PLEM_{slug(aeroporto)}.docx"
        salvar_docx(caminho, f"Fluxogramas PLEM — {aeroporto}", paginas)
        arquivos_individuais.append(caminho)
        print(
            f"{aeroporto}: {len(fluxos)} fluxos, "
            f"{sum(fluxo.quantidade_nos for fluxo in fluxos)} nós -> {caminho.name}"
        )

    salvar_consolidado_hibrido(
        ARQ_CONSOLIDADO,
        itens_consolidados,
        logo,
    )
    gerar_relatorio(por_aeroporto, sem_fluxo, orfaos, avisos)
    arquivos_manifesto = arquivos_individuais + [ARQ_CONSOLIDADO, ARQ_RELATORIO]
    gerar_manifesto(por_aeroporto, sem_fluxo, orfaos, avisos, arquivos_manifesto)
    gerar_pacote(arquivos_manifesto + [ARQ_MANIFESTO])

    print(f"Consolidado: {ARQ_CONSOLIDADO}")
    print(f"Relatório: {ARQ_RELATORIO}")
    print(f"Pacote: {ARQ_PACOTE}")
    print("TOTAL: 16 aeroportos, 129 fluxogramas, 1.976 nós.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
